#!/usr/bin/env python3
"""
Template for returns disposition audit: compares planned dispositions to event logs
using alias normalization and event status filtering.

Usage: python3 generate_returns_audit.py <plan.xlsx> <events.xlsx> <aliases.xlsx> <output.xlsx> <output.docx>

Adapt column indices, grouping keys, and qualifying status to match specific task requirements.
"""
import openpyxl
from docx import Document
from collections import defaultdict
import sys


def load_alias_map(wb):
    """Load alias sheet as case-insensitive lookup dict."""
    ws = wb.active
    alias_map = {}
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[0] is None or row[1] is None:
            continue
        alias_map[str(row[0]).strip().lower()] = str(row[1]).strip()
    return alias_map


def build_event_lookup(event_rows, qualifying_status='COMPLETED'):
    """
    Build lookup: (return_id, line_id) -> (event_time, final_disposition)
    Keeps only qualifying status events. Selects latest event per key.
    """
    lookup = {}
    for row in event_rows:
        ret_id, line_id, evt_time, evt_status, final_disp = row[0], row[1], row[2], row[3], row[4]
        if evt_status and str(evt_status).strip().upper() == qualifying_status.upper():
            key = (str(ret_id).strip(), str(line_id).strip())
            if key not in lookup or evt_time > lookup[key][0]:
                lookup[key] = (evt_time, str(final_disp).strip())
    return lookup


def compute_flags(plan_row, event_lookup, alias_map, planned_disp_idx=2):
    """
    Compute Missing Final Event, Disposition Mismatch, Total Errors, Error Summary.
    Returns dict of flag values.
    """
    ret_id, line_id = str(plan_row[0]).strip(), str(plan_row[1]).strip()
    key = (ret_id, line_id)
    planned_disp = str(plan_row[planned_disp_idx]).strip()

    if key not in event_lookup:
        missing = 1
        mismatch = 0
    else:
        missing = 0
        evt_disp = event_lookup[key][1]
        normalized = alias_map.get(evt_disp.lower(), evt_disp)
        mismatch = 0 if normalized.upper() == planned_disp.upper() else 1

    total = missing + mismatch
    errors = []
    if missing:
        errors.append('Missing Final Event')
    if mismatch:
        errors.append('Disposition Mismatch')
    summary = ', '.join(errors) or 'None'

    return {
        'Missing Final Event': missing,
        'Disposition Mismatch': mismatch,
        'Total Errors': total,
        'Error Summary': summary
    }


def main(plan_path, event_path, alias_path, excel_out, word_out):
    # Load inputs
    plan_wb = openpyxl.load_workbook(plan_path)
    event_wb = openpyxl.load_workbook(event_path)
    alias_wb = openpyxl.load_workbook(alias_path)

    plan_ws = plan_wb.active
    event_ws = event_wb.active

    plan_headers = [c.value for c in plan_ws[1]]
    plan_rows = list(plan_ws.iter_rows(min_row=2, values_only=True))
    event_rows = list(event_ws.iter_rows(min_row=2, values_only=True))

    alias_map = load_alias_map(alias_wb)
    event_lookup = build_event_lookup(event_rows)

    # Compute formatted data
    flag_names = ['Missing Final Event', 'Disposition Mismatch', 'Total Errors', 'Error Summary']
    fmt_headers = plan_headers + flag_names
    fmt_rows = []

    for row in plan_rows:
        flags = compute_flags(list(row), event_lookup, alias_map)
        fmt_row = list(row) + [flags[k] for k in flag_names]
        fmt_rows.append(fmt_row)

    # Aggregate summary by (Warehouse, Carrier) - adapt indices as needed
    wh_idx = plan_headers.index('Warehouse') if 'Warehouse' in plan_headers else 5
    carrier_idx = plan_headers.index('Carrier') if 'Carrier' in plan_headers else 6

    agg = defaultdict(lambda: [0, 0, 0])  # Missing, Mismatch, Total
    for row in fmt_rows:
        key = (row[wh_idx], row[carrier_idx])
        agg[key][0] += row[fmt_headers.index('Missing Final Event')]
        agg[key][1] += row[fmt_headers.index('Disposition Mismatch')]
        agg[key][2] += row[fmt_headers.index('Total Errors')]

    # Filter to error groups
    error_agg = {k: v for k, v in agg.items() if v[2] > 0}
    sorted_keys = sorted(error_agg.keys())

    summary_headers = ['Warehouse', 'Carrier', 'Missing Final Events', 'Disposition Mismatches', 'Total Errors']
    summary_rows = [[k[0], k[1], *v] for k, v in error_agg.items()]

    # Grand Total
    if summary_rows:
        totals = [sum(r[i] for r in summary_rows) for i in range(2, 5)]
        summary_rows.append(['Grand Total', '-', *totals])

    # Write Excel
    out_wb = openpyxl.Workbook()
    ws_raw = out_wb.active
    ws_raw.title = 'RawData'
    ws_raw.append(plan_headers)
    for r in plan_rows:
        ws_raw.append(r)

    ws_fmt = out_wb.create_sheet('Formatted Data')
    ws_fmt.append(fmt_headers)
    for r in fmt_rows:
        ws_fmt.append(r)

    ws_sum = out_wb.create_sheet('Summary')
    ws_sum.append(summary_headers)
    for r in summary_rows:
        ws_sum.append(r)
    out_wb.save(excel_out)

    # Write Word Brief
    doc = Document()
    doc.add_heading('Returns Disposition Audit - Executive Brief', level=1)
    doc.add_heading('Check Definitions', level=2)
    doc.add_paragraph('Missing Final Event: No COMPLETED event exists for the plan line.')
    doc.add_paragraph('Disposition Mismatch: Normalized final disposition does not match planned disposition.')
    doc.add_heading('Aggregate Totals', level=2)
    if summary_rows:
        total_errors = summary_rows[-1][-1] if isinstance(summary_rows[-1][-1], (int, float)) else 0
        doc.add_paragraph(f'Total Errors: {total_errors}')
    doc.add_heading('High-Priority Return IDs', level=2)
    doc.add_paragraph('Identify return IDs with frequent exceptions for targeted review.')
    doc.add_heading('Recommendations', level=2)
    doc.add_paragraph('Implement automated alerts for uncompleted dispositions within 24 hours.')
    doc.save(word_out)

    print(f'Outputs generated: {excel_out}, {word_out}')
    print(f'Summary: {len(summary_rows)-1} error groups, {summary_rows[-1][-1] if summary_rows else 0} total errors')


if __name__ == '__main__':
    if len(sys.argv) != 6:
        print('Usage: generate_returns_audit.py <plan.xlsx> <events.xlsx> <aliases.xlsx> <output.xlsx> <output.docx>')
        sys.exit(1)
    main(sys.argv[1], sys.argv[2], sys.argv[3], sys.argv[4], sys.argv[5])
