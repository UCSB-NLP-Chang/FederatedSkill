#!/usr/bin/env python3
"""
Template for generating Excel audit workbooks and Word executive briefs.
Adapt BUSINESS_RULES, GROUP_KEYS, and output paths as needed.
"""
import openpyxl
from docx import Document
from collections import defaultdict
import sys

def compute_flags(row, headers):
    """Return dict of derived flags based on row data. Adapt to actual column names."""
    expected = row[headers.index("Expected Qty")]
    received = row[headers.index("Received Qty")]
    storage = str(row[headers.index("Storage Class")]).strip().upper()
    temp = str(row[headers.index("Temp Status")]).strip().upper()
    
    qty_var = 1 if received != expected else 0
    cold_chain = 1 if storage in ("CHILLED", "FROZEN") and temp != "OK" else 0
    return {
        "Qty Variance": qty_var,
        "Cold Chain Error": cold_chain,
        "Total Errors": qty_var + cold_chain,
        "Error Summary": ", ".join(k for k, v in [("Qty Variance", qty_var), ("Cold Chain Error", cold_chain)] if v) or "None"
    }

def main(source_path, excel_out, word_out):
    wb = openpyxl.load_workbook(source_path)
    ws = wb.active
    rows = list(ws.iter_rows(min_row=1, max_row=ws.max_row, values_only=True))
    headers = list(rows[0])
    data = rows[1:]

    # Formatted Data
    fmt_headers = headers + ["Qty Variance", "Cold Chain Error", "Total Errors", "Error Summary"]
    fmt_rows = []
    for row in data:
        flags = compute_flags(list(row), headers)
        fmt_rows.append(list(row) + [flags["Qty Variance"], flags["Cold Chain Error"], flags["Total Errors"], flags["Error Summary"]])

    # Summary Aggregation
    group_keys = (headers.index("Item Code"), headers.index("Supplier"))
    agg = defaultdict(lambda: [0, 0, 0])
    for row in fmt_rows:
        key = (row[group_keys[0]], row[group_keys[1]])
        agg[key][0] += row[fmt_headers.index("Qty Variance")]
        agg[key][1] += row[fmt_headers.index("Cold Chain Error")]
        agg[key][2] += row[fmt_headers.index("Total Errors")]
    
    sorted_keys = sorted(agg.keys())
    summary_rows = [[k[0], k[1], v[0], v[1], v[2]] for k, v in agg.items()]
    totals = [sum(r[i] for r in summary_rows) for i in range(2, 5)]
    summary_rows.append(["Grand Total", "-", *totals])

    # Write Excel
    out_wb = openpyxl.Workbook()
    ws_raw = out_wb.active
    ws_raw.title = "RawData"
    ws_raw.append(headers)
    for r in data: ws_raw.append(r)

    ws_fmt = out_wb.create_sheet("Formatted Data")
    ws_fmt.append(fmt_headers)
    for r in fmt_rows: ws_fmt.append(r)

    ws_sum = out_wb.create_sheet("Summary")
    ws_sum.append(["Item Code", "Supplier", "Qty Variance Errors", "Cold Chain Errors", "Total Errors"])
    for r in summary_rows: ws_sum.append(r)
    out_wb.save(excel_out)

    # Write Word Brief
    doc = Document()
    doc.add_heading('Audit Executive Brief', level=1)
    doc.add_paragraph('Definitions and totals go here.')
    doc.add_paragraph(f'Total Errors: {totals[2]}')
    doc.save(word_out)
    print("Outputs generated successfully.")

if __name__ == "__main__":
    if len(sys.argv) != 4:
        print("Usage: generate_audit.py <source.xlsx> <output.xlsx> <output.docx>")
        sys.exit(1)
    main(sys.argv[1], sys.argv[2], sys.argv[3])