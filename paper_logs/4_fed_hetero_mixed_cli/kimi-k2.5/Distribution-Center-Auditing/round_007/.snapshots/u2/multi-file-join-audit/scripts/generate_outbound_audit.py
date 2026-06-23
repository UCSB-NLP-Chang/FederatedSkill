#!/usr/bin/env python3
"""
Template for outbound manifest audit: compares manifest plans to dock scan logs
to detect missing load scans and zone mismatches. Preserves template Overview sheet.

Usage: python3 generate_outbound_audit.py <template.xlsx> <manifest.xlsx> <scans.xlsx> <output.xlsx> <output.docx>

Adapt column indices, qualifying status, and grouping keys to match specific task requirements.
"""
import openpyxl
from docx import Document
from collections import defaultdict
import sys


def build_scan_lookup(scan_rows, ship_idx=0, carton_idx=1, zone_idx=2, time_idx=3, status_idx=4, qualifying='LOADED'):
    """
    Build lookup: (shipment_id, carton_id) -> (timestamp, scanned_zone)
    Keeps only qualifying status events. Selects latest event per key by timestamp.
    """
    lookup = {}
    for row in scan_rows:
        if row[status_idx] is None:
            continue
        status = str(row[status_idx]).strip().upper()
        if status == qualifying.upper():
            key = (str(row[ship_idx]).strip(), str(row[carton_idx]).strip())
            ts = row[time_idx]
            zone = str(row[zone_idx]).strip() if row[zone_idx] is not None else ''
            if key not in lookup or ts > lookup[key][0]:
                lookup[key] = (ts, zone)
    return lookup


def compute_flags(manifest_row, scan_lookup, ship_idx=0, carton_idx=1, zone_idx=2):
    """
    Compute Missing Load Scan, Zone Mismatch, Total Errors, Error Summary.
    Returns dict of flag values.
    """
    key = (str(manifest_row[ship_idx]).strip(), str(manifest_row[carton_idx]).strip())
    planned_zone = str(manifest_row[zone_idx]).strip() if manifest_row[zone_idx] is not None else ''
    
    if key not in scan_lookup:
        missing = 1
        mismatch = 0
    else:
        missing = 0
        scanned_zone = scan_lookup[key][1]
        mismatch = 0 if scanned_zone == planned_zone else 1
    
    total = missing + mismatch
    errors = []
    if missing:
        errors.append('Missing Load Scan')
    if mismatch:
        errors.append('Zone Mismatch')
    summary = ', '.join(errors) or 'None'
    
    return {
        'Missing Load Scan': missing,
        'Zone Mismatch': mismatch,
        'Total Errors': total,
        'Error Summary': summary
    }


def main(template_path, manifest_path, scans_path, excel_out, word_out):
    # Load inputs
    template_wb = openpyxl.load_workbook(template_path)
    manifest_wb = openpyxl.load_workbook(manifest_path)
    scans_wb = openpyxl.load_workbook(scans_path)
    
    manifest_ws = manifest_wb.active
    scans_ws = scans_wb.active
    
    manifest_headers = [c.value for c in manifest_ws[1]]
    manifest_rows = list(manifest_ws.iter_rows(min_row=2, values_only=True))
    scan_rows = list(scans_ws.iter_rows(min_row=2, values_only=True))
    
    # Build scan lookup (latest LOADED per carton)
    scan_lookup = build_scan_lookup(scan_rows)
    
    # Compute formatted data
    flag_names = ['Missing Load Scan', 'Zone Mismatch', 'Total Errors', 'Error Summary']
    fmt_headers = manifest_headers + flag_names
    fmt_rows = []
    
    for row in manifest_rows:
        flags = compute_flags(list(row), scan_lookup)
        fmt_row = list(row) + [flags[k] for k in flag_names]
        fmt_rows.append(fmt_row)
    
    # Aggregate summary by (Route, Shipment ID) - adapt indices as needed
    try:
        route_idx = manifest_headers.index('Route')
        ship_idx = manifest_headers.index('Shipment ID')
    except ValueError:
        route_idx = 3  # fallback
        ship_idx = 0   # fallback
    
    agg = defaultdict(lambda: [0, 0, 0])  # Missing, Mismatch, Total
    for row in fmt_rows:
        key = (row[route_idx], row[ship_idx])
        agg[key][0] += row[fmt_headers.index('Missing Load Scan')]
        agg[key][1] += row[fmt_headers.index('Zone Mismatch')]
        agg[key][2] += row[fmt_headers.index('Total Errors')]
    
    # Filter to error groups only
    error_agg = {k: v for k, v in agg.items() if v[2] > 0}
    sorted_keys = sorted(error_agg.keys())
    
    summary_headers = ['Route', 'Shipment ID', 'Missing Load Scans', 'Zone Mismatches', 'Total Errors']
    summary_rows = [[k[0], k[1], *v] for k, v in error_agg.items()]
    
    # Grand Total
    if summary_rows:
        totals = [sum(r[i] for r in summary_rows) for i in range(2, 5)]
        summary_rows.append(['Grand Total', '-', *totals])
    
    # Create output workbook, preserving template Overview
    out_wb = openpyxl.Workbook()
    
    # Copy Overview from template
    if 'Overview' in template_wb.sheetnames:
        src = template_wb['Overview']
        dst = out_wb.active
        dst.title = 'Overview'
        for row in src.iter_rows(values_only=True):
            dst.append(row)
    
    # Remove default sheet if it exists and we added Overview
    if 'Sheet' in out_wb.sheetnames:
        out_wb.remove(out_wb['Sheet'])
    
    # RawData
    ws_raw = out_wb.create_sheet('RawData')
    ws_raw.append(manifest_headers)
    for r in manifest_rows:
        ws_raw.append(r)
    
    # Formatted Data
    ws_fmt = out_wb.create_sheet('Formatted Data')
    ws_fmt.append(fmt_headers)
    for r in fmt_rows:
        ws_fmt.append(r)
    
    # Summary
    ws_sum = out_wb.create_sheet('Summary')
    ws_sum.append(summary_headers)
    for r in summary_rows:
        ws_sum.append(r)
    
    out_wb.save(excel_out)
    
    # Write Word Brief
    doc = Document()
    doc.add_heading('Outbound Load Audit - Executive Summary', level=1)
    doc.add_paragraph(
        'This audit assessed carton handoff accuracy by comparing the manifest plan against dock scan events. '
        'Two compliance checks were performed:'
    )
    doc.add_paragraph(
        '• Missing Load Scan: Indicates when a carton listed in the manifest has no corresponding LOADED scan record '
        'in the dock log, suggesting a missed handoff.'
    )
    doc.add_paragraph(
        '• Zone Mismatch: Indicates when a carton was loaded from a different zone than planned in the manifest, '
        'suggesting a potential misrouting or incorrect staging area usage.'
    )
    
    if summary_rows:
        total_missing = summary_rows[-1][2]
        total_mismatch = summary_rows[-1][3]
        total_errors = summary_rows[-1][4]
        doc.add_paragraph(
            f'The audit identified {total_missing} Missing Load Scans, {total_mismatch} Zone Mismatches, '
            f'and {total_errors} Total Errors across the dataset.'
        )
        
        # Identify high-priority shipments (more than 1 error)
        high_priority = [(k, v) for k, v in error_agg.items() if v[2] > 1]
        if high_priority:
            doc.add_heading('High-Priority Shipments', level=2)
            doc.add_paragraph('Shipments requiring immediate attention:')
            for (route, ship_id), counts in sorted(high_priority):
                doc.add_paragraph(f'• {ship_id} (Route: {route}): {counts[2]} errors ({counts[0]} Missing, {counts[1]} Mismatch)')
    
    doc.add_heading('Recommendations', level=2)
    doc.add_paragraph(
        '1. Verify scanner availability and operator training at dock doors handling high-error routes.'
    )
    doc.add_paragraph(
        '2. Implement zone verification checkpoints before cartons are released to carriers.'
    )
    doc.add_paragraph(
        '3. Review missing scan procedures for flagged shipments to ensure all cartons are properly logged before departure.'
    )
    
    doc.save(word_out)
    print(f'Outputs generated: {excel_out}, {word_out}')
    if summary_rows:
        print(f'Summary: {len(summary_rows)-1} error groups, {summary_rows[-1][-1]} total errors')


if __name__ == '__main__':
    if len(sys.argv) != 6:
        print('Usage: generate_outbound_audit.py <template.xlsx> <manifest.xlsx> <scans.xlsx> <output.xlsx> <output.docx>')
        sys.exit(1)
    main(sys.argv[1], sys.argv[2], sys.argv[3], sys.argv[4], sys.argv[5])
