#!/usr/bin/env python3
"""
Template for generating Excel audit workbooks and Word executive briefs.
Adapt BUSINESS_RULES, GROUP_KEYS, and output paths as needed.
Supports multi-sheet sources with reference data lookups.
"""
import openpyxl
from docx import Document
from collections import defaultdict
import sys

def load_reference_rules(wb, sheet_name):
    """
    Load a reference sheet as a lookup dictionary.
    Assumes first column is the key. Adapt as needed.
    Returns: {key: {col2: val, col3: val, ...}}
    """
    if sheet_name not in wb.sheetnames:
        return {}
    ws = wb[sheet_name]
    rows = list(ws.iter_rows(min_row=1, values_only=True))
    if not rows:
        return {}
    headers = list(rows[0])
    rules = {}
    for row in rows[1:]:
        if row[0] is None:
            continue
        key = row[0]
        rules[key] = {headers[i]: row[i] for i in range(1, len(headers))}
    return rules

def compute_flags(row, headers, sla_rules=None):
    """
    Return dict of derived flags based on row data.
    Adapt to actual column names and business rules.
    
    Example for receipt audit (direct comparison):
    - expected = row[headers.index("Expected Qty")]
    - received = row[headers.index("Received Qty")]
    - qty_var = 1 if received != expected else 0
    
    Example for SLA audit (reference lookup):
    - priority = row[headers.index("Priority Tier")]
    - open_age = row[headers.index("Open Age Hours")]
    - max_hours = sla_rules.get(priority, {}).get("Max Open Hours", 999)
    - sla_breach = 1 if open_age > max_hours else 0
    """
    # Default implementation - replace with actual logic
    expected = row[headers.index("Expected Qty")]
    received = row[headers.index("Received Qty")]
    storage = str(row[headers.index("Storage Class")]).strip().upper()
    temp = str(row[headers.index("Temp Status")]).strip().upper()
    
    qty_var = 1 if received != expected else 0
    cold_chain = 1 if storage in ("CHILLED", "FROZEN") and temp != "OK" else 0
    return {
        "Qty Variance": qty_var,
        "Cold Chain Error": cold_chain,
        "Total Errors": qty_var + cold_chain,
        "Error Summary": ", ".join(k for k, v in [("Qty Variance", qty_var), ("Cold Chain Error", cold_chain)] if v) or "None"
    }

def main(source_path, excel_out, word_out):
    wb = openpyxl.load_workbook(source_path)
    
    # Load primary data from active sheet or named sheet
    ws = wb.active
    if 'Tickets' in wb.sheetnames:  # Common for SLA audits
        ws = wb['Tickets']
    elif 'Data' in wb.sheetnames:
        ws = wb['Data']
    
    rows = list(ws.iter_rows(min_row=1, max_row=ws.max_row, values_only=True))
    headers = list(rows[0])
    data = rows[1:]

    # Load reference rules if available (e.g., SLA_Rules sheet)
    sla_rules = load_reference_rules(wb, 'SLA_Rules') if 'SLA_Rules' in wb.sheetnames else {}

    # Formatted Data - add computed columns
    # Adapt flag_names based on your audit type
    flag_names = ["Qty Variance", "Cold Chain Error", "Total Errors", "Error Summary"]
    fmt_headers = headers + flag_names
    fmt_rows = []
    
    for row in data:
        flags = compute_flags(list(row), headers, sla_rules)
        fmt_row = list(row) + [flags.get(k, 0 if 'Error' in k else 'None') for k in flag_names]
        fmt_rows.append(fmt_row)

    # Summary Aggregation
    # Adapt group_keys indices based on your data structure
    # Example: group by Item Code and Supplier
    group_indices = [headers.index("Item Code"), headers.index("Supplier")] if "Item Code" in headers else [0]
    
    agg = defaultdict(lambda: [0] * len(flag_names[:-1]))  # Exclude Error Summary from sums
    for row in fmt_rows:
        key = tuple(row[i] for i in group_indices)
        for i, flag in enumerate(flag_names[:-1]):  # Skip Error Summary
            val = row[fmt_headers.index(flag)]
            agg[key][i] += val if isinstance(val, (int, float)) else 0
    
    # Filter to error rows only (optional - remove if all groups needed)
    error_agg = {k: v for k, v in agg.items() if sum(v) > 0}
    
    sorted_keys = sorted(error_agg.keys())
    summary_headers = [headers[i] for i in group_indices] + flag_names[:-1]
    summary_rows = [[*k, *v] for k, v in error_agg.items()]
    
    # Grand Total
    if summary_rows:
        totals = [sum(r[len(group_indices) + i] for r in summary_rows) for i in range(len(flag_names[:-1]))]
        summary_rows.append(["Grand Total"] + ["-"] * (len(group_indices) - 1) + totals)

    # Write Excel
    out_wb = openpyxl.Workbook()
    
    ws_raw = out_wb.active
    ws_raw.title = "RawData"
    ws_raw.append(headers)
    for r in data:
        ws_raw.append(r)

    ws_fmt = out_wb.create_sheet("Formatted Data")
    ws_fmt.append(fmt_headers)
    for r in fmt_rows:
        ws_fmt.append(r)

    ws_sum = out_wb.create_sheet("Summary")
    ws_sum.append(summary_headers)
    for r in summary_rows:
        ws_sum.append(r)
    out_wb.save(excel_out)

    # Write Word Brief
    doc = Document()
    doc.add_heading('Audit Executive Brief', level=1)
    doc.add_paragraph('Definitions and totals go here.')
    if summary_rows and len(summary_rows[-1]) > len(group_indices):
        total_errors = summary_rows[-1][-1] if isinstance(summary_rows[-1][-1], (int, float)) else 0
        doc.add_paragraph(f'Total Errors: {total_errors}')
    doc.save(word_out)
    print("Outputs generated successfully.")
    print(f"Summary: {len(summary_rows)-1} error groups, {sum(summary_rows[-1][len(group_indices):]) if summary_rows else 0} total errors")

if __name__ == "__main__":
    if len(sys.argv) != 4:
        print("Usage: generate_audit.py <source.xlsx> <output.xlsx> <output.docx>")
        sys.exit(1)
    main(sys.argv[1], sys.argv[2], sys.argv[3])