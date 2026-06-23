#!/usr/bin/env python3
"""Template for creating multi-sheet Excel audit reports with Word summaries."""

import pandas as pd
from openpyxl import Workbook, load_workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from docx import Document


def create_audit_report(
    source_path: str,
    output_excel: str,
    output_docx: str = None,
    error_columns: dict = None,
    groupby_cols: list = None,
    summary_title: str = "Audit Summary"
):
    """
    Create a multi-sheet Excel audit report with optional Word summary.

    Args:
        source_path: Path to source Excel file
        output_excel: Path for output Excel workbook
        output_docx: Optional path for Word summary document
        error_columns: Dict mapping column names to (condition_func, description) tuples
        groupby_cols: Columns to group by for summary sheet
        summary_title: Title for Word document
    """
    # Read source data
    df = pd.read_excel(source_path)
    df_raw = df.copy()

    if groupby_cols is None:
        groupby_cols = ['Item Code', 'Supplier']

    # Create workbook
    wb = Workbook()

    # RawData sheet
    ws_raw = wb.active
    ws_raw.title = "RawData"
    for r_idx, row in enumerate(dataframe_to_rows(df_raw, index=False, header=True), 1):
        for c_idx, value in enumerate(row, 1):
            ws_raw.cell(row=r_idx, column=c_idx, value=value)

    # Formatted Data sheet with calculated columns
    df_formatted = df.copy()
    if error_columns:
        for col_name, (condition_func, _) in error_columns.items():
            df_formatted[col_name] = df.apply(condition_func, axis=1).astype(int)

        error_col_names = list(error_columns.keys())
        df_formatted['Total Errors'] = df_formatted[error_col_names].sum(axis=1)

        def build_error_summary(row):
            errors = []
            for col_name, (_, desc) in error_columns.items():
                if row[col_name] == 1:
                    errors.append(desc)
            return ', '.join(errors) if errors else 'None'
        df_formatted['Error Summary'] = df_formatted.apply(build_error_summary, axis=1)

    ws_formatted = wb.create_sheet("Formatted Data")
    for r_idx, row in enumerate(dataframe_to_rows(df_formatted, index=False, header=True), 1):
        for c_idx, value in enumerate(row, 1):
            ws_formatted.cell(row=r_idx, column=c_idx, value=value)

    # Summary sheet
    if error_columns:
        error_col_names = list(error_columns.keys())
        summary = df_formatted.groupby(groupby_cols, as_index=False)[error_col_names].sum()
        summary['Total Errors'] = summary[error_col_names].sum(axis=1)
        summary = summary[summary['Total Errors'] > 0]
        summary = summary.sort_values(by=groupby_cols[0])

        grand_total = {col: '-' if col in groupby_cols else summary[col].sum()
                       for col in summary.columns}
        summary = pd.concat([summary, pd.DataFrame([grand_total])], ignore_index=True)

        ws_summary = wb.create_sheet("Summary")
        for r_idx, row in enumerate(dataframe_to_rows(summary, index=False, header=True), 1):
            for c_idx, value in enumerate(row, 1):
                ws_summary.cell(row=r_idx, column=c_idx, value=value)

    wb.save(output_excel)

    # Verify "None" strings were written correctly
    wb_check = load_workbook(output_excel)
    ws_check = wb_check['Formatted Data']
    error_summary_col = None
    for idx, cell in enumerate(ws_check[1], 1):
        if cell.value == 'Error Summary':
            error_summary_col = idx
            break

    if error_summary_col:
        for row in ws_check.iter_rows(min_row=2, max_row=ws_check.max_row):
            val = row[error_summary_col - 1].value
            if val is None or str(val) == 'nan':
                row[error_summary_col - 1].value = 'None'
        wb_check.save(output_excel)

    # Word summary (optional)
    if output_docx:
        doc = Document()
        doc.add_heading(summary_title, level=1)

        total_errors = int(df_formatted['Total Errors'].sum()) if 'Total Errors' in df_formatted.columns else 0
        para = doc.add_paragraph()
        para.add_run(f"Total exceptions identified: {total_errors}")

        doc.save(output_docx)

    return df_formatted


if __name__ == "__main__":
    print("Use create_audit_report() function with appropriate parameters")
