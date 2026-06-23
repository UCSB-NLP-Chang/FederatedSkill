#!/usr/bin/env python3
"""
Generate Excel audit workbook with RawData, Formatted Data, Summary sheets
plus Word executive brief.

Usage:
    python3 generate_audit_report.py <input.xlsx> <output_prefix> [options]

Outputs:
    <output_prefix>_Audit.xlsx
    <output_prefix>_Brief.docx
"""

import sys
import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils.dataframe import dataframe_to_rows
from docx import Document
from docx.shared import Inches


def calculate_exceptions(df, exception_rules, summary_col_name='Error Summary'):
    """
    Calculate exception columns based on configurable rules.
    
    Args:
        df: Source DataFrame
        exception_rules: List of dicts with keys:
            - 'name': column name for the exception
            - 'condition': boolean Series indicating error rows
        summary_col_name: Name for the human-readable summary column
    
    Returns:
        DataFrame with added exception columns + Total Errors + Error Summary
    """
    result = df.copy()
    
    exception_names = []
    for rule in exception_rules:
        name = rule['name']
        condition = rule['condition']
        result[name] = condition.astype(int)
        exception_names.append(name)
    
    # Total errors
    result['Total Errors'] = result[exception_names].sum(axis=1)
    
    # Human-readable summary
    def make_summary(row):
        parts = [name for name in exception_names if row[name]]
        return ', '.join(parts) if parts else 'None'
    
    result[summary_col_name] = result.apply(make_summary, axis=1)
    
    return result


def create_summary(df, group_cols, error_cols, sort_cols=None):
    """Create aggregated summary with Grand Total row."""
    if sort_cols is None:
        sort_cols = group_cols
    
    agg_dict = {col: 'sum' for col in error_cols}
    summary = df.groupby(group_cols).agg(agg_dict).reset_index()
    summary = summary.sort_values(sort_cols)
    
    # Add Grand Total
    totals = summary[error_cols].sum()
    grand_total = {col: '-' for col in group_cols}
    grand_total.update({col: totals[col] for col in error_cols})
    grand_total[group_cols[0]] = 'Grand Total'
    
    summary = pd.concat([summary, pd.DataFrame([grand_total])], ignore_index=True)
    return summary


def write_excel_audit(df_raw, df_formatted, df_summary, output_path,
                      sheet_names=None, header_fill='4472C4', header_font='FFFFFF'):
    """Write multi-sheet Excel audit workbook with formatting."""
    if sheet_names is None:
        sheet_names = ['RawData', 'Formatted Data', 'Summary']
    
    wb = Workbook()
    wb.remove(wb.active)
    
    sheets_data = [
        (sheet_names[0], df_raw),
        (sheet_names[1], df_formatted),
        (sheet_names[2], df_summary)
    ]
    
    for sheet_name, data in sheets_data:
        ws = wb.create_sheet(title=sheet_name)
        
        for r_idx, row in enumerate(dataframe_to_rows(data, index=False, header=True), 1):
            for c_idx, value in enumerate(row, 1):
                cell = ws.cell(row=r_idx, column=c_idx, value=value)
                
                if r_idx == 1:
                    cell.font = Font(bold=True, color=header_font)
                    cell.fill = PatternFill(start_color=header_fill, end_color=header_fill, fill_type='solid')
                    cell.alignment = Alignment(horizontal='center')
    
    wb.save(output_path)
    return output_path


def write_word_brief(output_path, title, definitions, totals, 
                     high_priority_items, recommendations, context_desc=""):
    """
    Generate executive brief Word document.
    
    Args:
        output_path: Path for .docx file
        title: Document title
        definitions: Dict of {term: description}
        totals: Dict of {metric: value}
        high_priority_items: List of item identifiers (e.g., ['CarrierA', 'CarrierE'])
        recommendations: List of recommendation strings
        context_desc: Brief description of audit scope (e.g., "16 loads from 5 carriers")
    """
    doc = Document()
    doc.add_heading(title, 0)
    
    # Context paragraph
    if context_desc:
        doc.add_paragraph(f"This audit assessed {context_desc}.")
    
    # Definitions
    for term, desc in definitions.items():
        doc.add_paragraph(f"{term}: {desc}")
    
    doc.add_paragraph()
    
    # Totals
    totals_para = doc.add_paragraph()
    totals_text = "The audit identified "
    total_items = list(totals.items())
    for i, (metric, value) in enumerate(total_items):
        if i == len(total_items) - 1 and len(total_items) > 1:
            totals_text += f"and {value} {metric}"
        else:
            totals_text += f"{value} {metric}"
            if i < len(total_items) - 2:
                totals_text += ", "
            elif i == len(total_items) - 2 and len(total_items) == 2:
                totals_text += " "
            elif i < len(total_items) - 1:
                totals_text += ", "
    totals_text += "."
    totals_para.add_run(totals_text)
    
    # High priority
    if high_priority_items:
        doc.add_paragraph()
        if len(high_priority_items) == 1:
            priority_text = f"Priority attention is recommended for {high_priority_items[0]}."
        else:
            priority_text = f"Priority attention is recommended for {', '.join(high_priority_items[:-1])} and {high_priority_items[-1]}."
        doc.add_paragraph(priority_text)
    
    # Recommendations
    if recommendations:
        doc.add_paragraph()
        rec_para = doc.add_paragraph("Recommendations: ")
        for i, rec in enumerate(recommendations):
            if i == len(recommendations) - 1 and len(recommendations) > 1:
                rec_para.add_run(f"and {rec}")
            else:
                rec_para.add_run(rec)
                if i < len(recommendations) - 2:
                    rec_para.add_run(", ")
                elif i < len(recommendations) - 1:
                    rec_para.add_run(" ")
        rec_para.add_run(".")
    
    doc.save(output_path)
    return output_path


def main():
    if len(sys.argv) < 3:
        print(f"Usage: {sys.argv[0]} <input.xlsx> <output_prefix>")
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_prefix = sys.argv[2]
    
    # Read source
    df_raw = pd.read_excel(input_path)
    print(f"Source data loaded: {len(df_raw)} rows")
    
    # === CONFIGURE EXCEPTION RULES FOR YOUR AUDIT ===
    # Example: Trailer detention audit
    exception_rules = [
        {
            'name': 'Detention Overrun',
            'condition': df_raw['Actual Hold Hours'] > df_raw['Allowed Hold Hours']
        },
        {
            'name': 'Seal Error',
            'condition': (df_raw['Seal Required'] == 'YES') & 
                        (~df_raw['Seal Status'].isin(['VERIFIED']))
        }
    ]
    
    group_cols = ['Carrier', 'Yard']
    error_cols = ['Detention Overrun', 'Seal Error', 'Total Errors']
    
    definitions = {
        "Detention Overrun": "occurs when actual hold hours exceed allowed hold hours",
        "Seal Error": "flagged when a load requires a seal but status is not VERIFIED"
    }
    
    recommendations = [
        "review detention policies with carriers showing repeated overruns",
        "enforce stricter seal verification protocols at all yard locations"
    ]
    
    context_desc = f"{len(df_raw)} loads from {df_raw['Carrier'].nunique()} carriers"
    # === END CONFIGURATION ===
    
    # Calculate exceptions
    df_formatted = calculate_exceptions(df_raw, exception_rules)
    print(f"Formatted data created with {len(exception_rules)} exception types")
    
    # Create summary
    df_summary = create_summary(df_formatted, group_cols, error_cols)
    print(f"Summary created: {len(df_summary)-1} groups + Grand Total")
    
    # Write Excel
    excel_path = f"{output_prefix}_Audit.xlsx"
    write_excel_audit(df_raw, df_formatted, df_summary, excel_path)
    print(f"Excel file created: {excel_path}")
    
    # Identify high priority (top 2 by total errors)
    priority_df = df_summary[df_summary[group_cols[0]] != 'Grand Total'].nlargest(2, 'Total Errors')
    high_priority = priority_df[group_cols[0]].tolist()
    
    # Calculate totals for brief
    totals = {
        "detention overrun errors": int(df_formatted['Detention Overrun'].sum()),
        "seal compliance errors": int(df_formatted['Seal Error'].sum()),
        "total exceptions": int(df_formatted['Total Errors'].sum())
    }
    
    # Write Word brief
    doc_path = f"{output_prefix}_Brief.docx"
    write_word_brief(
        doc_path,
        title="Compliance Audit - Executive Summary",
        definitions=definitions,
        totals=totals,
        high_priority_items=high_priority,
        recommendations=recommendations,
        context_desc=context_desc
    )
    print(f"Word file created: {doc_path}")
    
    print(f"\n=== Audit Generation Complete ===")
    for name, val in totals.items():
        print(f"{name.replace('_', ' ').title()}: {val}")
    print(f"High Priority: {', '.join(high_priority)}")


if __name__ == '__main__':
    main()
